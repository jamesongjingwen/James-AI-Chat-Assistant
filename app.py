from flask import Flask, render_template, request, jsonify
import os
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from docx import Document
import warnings

warnings.filterwarnings("ignore")

# Set the API key as an environment variable
api_key = '' #Input your API Key here
os.environ["OPENAI_API_KEY"] = api_key

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    full_text_string = '\n'.join(full_text)
    return full_text_string

file_path = r'C:\Users\theja\OneDrive\Desktop\MSBA\MSBA Course\Semester 2\SAP Test\VS\Summary.docx' #Enter your Summary.docx dataset directory here
data = read_docx(file_path)

class SimpleCharacterTextSplitter:
    def __init__(self, chunk_size, chunk_overlap):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_documents(self, document):
        chunks = []
        start = 0
        while start < len(document):
            end = start + self.chunk_size
            if end < len(document):  
                while end > start and document[end] not in [' ', '\n', '.', ',']:
                    end -= 1
            chunks.append(document[start:end])
            start = end if end + self.chunk_overlap > len(document) else end - self.chunk_overlap
        return chunks

text_splitter = SimpleCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
split_data = text_splitter.split_documents(data)

embeddings = OpenAIEmbeddings()
class Document:
    def __init__(self, page_content, metadata={}):  # Use an empty dict as default
        self.page_content = page_content
        self.metadata = metadata
transformed_data = [Document(doc, metadata={"key": "value"}) for doc in split_data]

vectorstore = FAISS.from_documents(transformed_data, embedding=embeddings)

llm = ChatOpenAI(temperature=0.7, model_name="gpt-4")

memory = ConversationBufferMemory(memory_key='chat_history', return_messages=True)

conversation_chain = ConversationalRetrievalChain.from_llm(
    llm=llm,
    chain_type="stuff",
    retriever=vectorstore.as_retriever(),
    memory=memory
)

app = Flask(__name__)

@app.route('/')
def home():
    return render_template('chat.html')

@app.route('/ask', methods=['POST'])
def ask():
    try:
        user_message = request.form['message']
        result = conversation_chain({"question": user_message})
        response = result["answer"]  # Extract the answer from the result
    except Exception as e:
        response = "Sorry, I couldn't process your request."
        print(f"Error processing request: {e}")
    return jsonify({"response": response})

if __name__ == '__main__':
    app.run(debug=True)

